# Parses PDF and DOCX into a list of text blocks with a hint of which lines
# look like headings. DOCX is easy (paragraph styles tell us directly). For
# PDF we look at font sizes per line and call lines bigger than the median a
# "heading".
from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from typing import Optional

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


@dataclass
class TextBlock:
    text: str
    is_heading: bool = False
    style: Optional[str] = None
    font_size: Optional[float] = None


@dataclass
class ParsedDocument:
    raw_text: str
    blocks: list[TextBlock] = field(default_factory=list)


def parse(file_bytes: bytes, filename: str) -> ParsedDocument:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _parse_pdf(file_bytes)
    if name.endswith(".docx"):
        return _parse_docx(file_bytes)
    if name.endswith(".doc"):
        raise ValueError("Legacy .doc files are not supported. Please upload PDF or DOCX.")
    if name.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")
        blocks = [TextBlock(text=line) for line in text.splitlines() if line.strip()]
        return ParsedDocument(raw_text=text, blocks=blocks)
    raise ValueError(f"Unsupported file type: {filename}")


def _parse_docx(file_bytes: bytes) -> ParsedDocument:
    doc = Document(io.BytesIO(file_bytes))
    blocks: list[TextBlock] = []
    text_parts: list[str] = []
    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if not txt:
            continue
        style = (para.style.name if para.style else "") or ""
        is_heading = style.lower().startswith("heading") or style.lower() == "title"
        blocks.append(TextBlock(text=txt, is_heading=is_heading, style=style))
        text_parts.append(txt)
    return ParsedDocument(raw_text="\n".join(text_parts), blocks=blocks)


def _parse_pdf(file_bytes: bytes) -> ParsedDocument:
    blocks: list[TextBlock] = []
    text_parts: list[str] = []
    all_sizes: list[float] = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        # pass 1: gather font sizes so we know what counts as "big"
        for page in pdf.pages:
            for ch in page.chars or []:
                size = ch.get("size")
                if size:
                    all_sizes.append(float(size))

        median_size = _median(all_sizes) if all_sizes else 10.0
        heading_threshold = median_size + 1.5

        for page in pdf.pages:
            try:
                lines = _group_lines(page)
            except Exception as e:
                logger.warning("Falling back to plain extract_text() for a page: %s", e)
                txt = page.extract_text() or ""
                for line in txt.splitlines():
                    if line.strip():
                        blocks.append(TextBlock(text=line.strip()))
                        text_parts.append(line.strip())
                continue

            for line_text, max_size in lines:
                if not line_text.strip():
                    continue
                is_heading = max_size >= heading_threshold and len(line_text) < 80
                blocks.append(TextBlock(text=line_text.strip(), is_heading=is_heading, font_size=max_size))
                text_parts.append(line_text.strip())

    return ParsedDocument(raw_text="\n".join(text_parts), blocks=blocks)


def _group_lines(page) -> list[tuple[str, float]]:
    # Group chars by their y position, then read left to right.
    # Returns (line_text, biggest font on that line).
    chars = page.chars or []
    if not chars:
        txt = page.extract_text() or ""
        return [(line, 10.0) for line in txt.splitlines()]

    chars_sorted = sorted(chars, key=lambda c: (round(c["top"], 1), c["x0"]))
    lines: list[tuple[str, float]] = []
    current_top = None
    current_chars: list[dict] = []
    tolerance = 2.0

    def flush():
        if not current_chars:
            return
        current_chars.sort(key=lambda c: c["x0"])
        text = ""
        prev_x1 = None
        for c in current_chars:
            if prev_x1 is not None and (c["x0"] - prev_x1) > 1.5:
                text += " "
            text += c["text"]
            prev_x1 = c["x1"]
        max_size = max((c.get("size", 0) for c in current_chars), default=0.0)
        lines.append((text, float(max_size)))

    for c in chars_sorted:
        top = c["top"]
        if current_top is None or abs(top - current_top) <= tolerance:
            current_chars.append(c)
            current_top = top if current_top is None else current_top
        else:
            flush()
            current_chars = [c]
            current_top = top
    flush()
    return lines


def _median(values: list[float]) -> float:
    if not values:
        return 0.0
    s = sorted(values)
    n = len(s)
    return s[n // 2] if n % 2 else (s[n // 2 - 1] + s[n // 2]) / 2
